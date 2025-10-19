"""
Template processing module for mail merge application.
Handles template validation, field extraction, and placeholder processing.
"""

import re
from pathlib import Path
from docx import Document
from typing import List, Tuple, Dict, Any


def extract_placeholders(docx_file) -> Tuple[List[str], List[str]]:
    """
    Extract placeholders from a Word document.
    
    Args:
        docx_file: Uploaded file object or file path
        
    Returns:
        Tuple of (curly_fields, square_fields) - lists of field names
    """
    doc = Document(docx_file)
    curly_fields = []
    square_fields = []
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        curly_fields += [field.strip() for field in re.findall(r"\{\{(.*?)\}\}", para.text)]
        square_fields += re.findall(r"\[(.*?)\]", para.text)
    
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    curly_fields += [field.strip() for field in re.findall(r"\{\{(.*?)\}\}", para.text)]
                    square_fields += re.findall(r"\[(.*?)\]", para.text)
    
    # Extract from headers and footers
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            for para in hf.paragraphs:
                curly_fields += [field.strip() for field in re.findall(r"\{\{(.*?)\}\}", para.text)]
                square_fields += re.findall(r"\[(.*?)\]", para.text)
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            curly_fields += [field.strip() for field in re.findall(r"\{\{(.*?)\}\}", para.text)]
                            square_fields += re.findall(r"\[(.*?)\]", para.text)
    
    # Deduplicate while preserving order
    def _dedup(seq):
        seen = set()
        out = []
        for x in seq:
            if x not in seen:
                seen.add(x)
                out.append(x)
        return out

    curly_fields = _dedup(curly_fields)
    square_fields = _dedup(square_fields)
    return curly_fields, square_fields


def normalize(text: str) -> str:
    """Normalize text for matching purposes."""
    return re.sub(r"[\s_]+", "", text.strip().lower())


def is_valid_jinja_var(name: str) -> bool:
    """
    Check if a variable name is valid for Jinja templates.
    Supports both simple variables and complex expressions.
    
    Args:
        name: Variable name or expression to validate
        
    Returns:
        True if valid Jinja identifier or expression
    """
    name = name.strip()
    
    # First check if it's a simple variable name
    if re.match(r'^[A-Za-z_][A-Za-z0-9_]*(\.[A-Za-z_][A-Za-z0-9_]*)*$', name):
        return True
    
    # Check for complex expressions with operators, filters, conditionals, etc.
    # Allow expressions with:
    # - String concatenation (~)
    # - Filters (|)
    # - Conditionals (if/else)
    # - Parentheses for grouping
    # - Quotes for strings
    # - Brackets for lists
    # - Default values
    
    # Basic pattern for complex Jinja expressions
    complex_pattern = r'^[A-Za-z0-9_\s\.\[\]\(\)\'\"~|,:-]+$'
    
    # Check if it contains Jinja-specific operators/filters
    jinja_indicators = ['|', '~', 'if ', 'else', 'default(', 'trim', 'join(', 'reject(', 'equalto']
    
    if any(indicator in name for indicator in jinja_indicators):
        # Additional validation for complex expressions
        # Check for balanced parentheses and quotes
        paren_count = name.count('(') - name.count(')')
        if paren_count != 0:
            return False
            
        # Check for balanced quotes
        single_quotes = name.count("'")
        double_quotes = name.count('"')
        if single_quotes % 2 != 0 or double_quotes % 2 != 0:
            return False
            
        return True
    
    # Check for list expressions
    if name.startswith('[') and name.endswith(']'):
        return True
    
    # Check for string concatenation expressions
    if '~' in name and not re.search(r'[^A-Za-z0-9_\s\.\[\]\(\)\'\"~|,:-]', name):
        return True
    
    return False


def detect_macros(docx_file) -> bool:
    """
    Detect if a template uses Jinja macros.
    
    Args:
        docx_file: Uploaded file object or file path
        
    Returns:
        True if template contains macros
    """
    doc = Document(docx_file)
    
    # Check for macro definitions in paragraphs
    for para in doc.paragraphs:
        if re.search(r'\{%\s*macro\s+\w+', para.text):
            return True
    
    # Check for macro definitions in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if re.search(r'\{%\s*macro\s+\w+', para.text):
                        return True
    
    # Check for macro definitions in headers and footers
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            for para in hf.paragraphs:
                if re.search(r'\{%\s*macro\s+\w+', para.text):
                    return True
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if re.search(r'\{%\s*macro\s+\w+', para.text):
                                return True
    
    return False


def validate_template_placeholders(curly_fields: List[str]) -> Dict[str, List[str]]:
    """
    Validate template placeholders for common Jinja/docxtpl errors.
    
    Args:
        curly_fields: List of curly bracket field names
        
    Returns:
        Dictionary with validation results
    """
    validation_results = {
        'invalid': [],
        'control_in_print': [],
        'unclosed': []
    }
    
    for field in curly_fields:
        if not is_valid_jinja_var(field):
            validation_results['invalid'].append(field)
        if re.search(r"\b(for|if|else|endif|endfor)\b", field):
            validation_results['control_in_print'].append(field)
        if field.strip().lower().startswith("note:") or field.strip().endswith("if"):
            validation_results['unclosed'].append(field)
    
    return validation_results


def render_text(s: str) -> str:
    """Placeholder for any extra formatting/escaping in previews."""
    return s


def clean_placeholder_elements(doc):
    """
    Remove empty paragraphs and table rows that were placeholders from the document.
    
    Args:
        doc: DocxTemplate document object
    """
    # Clean empty paragraphs - more aggressive approach
    paragraphs_to_remove = []
    for i, p in enumerate(doc.docx.paragraphs):
        # Remove paragraphs that are empty or contain only Jinja code remnants
        if not p.text.strip():
            paragraphs_to_remove.append(i)
        elif re.search(r"^\s*\{[#%].*?[#%]\}\s*$", p.text.strip()):
            # Remove Jinja comments and control blocks that are now empty
            paragraphs_to_remove.append(i)
        elif re.search(r"^\s*\{%\s*(macro|set|if|for|endmacro|endif|endfor).*?%\}\s*$", p.text.strip()):
            # Remove macro definitions and control structures
            paragraphs_to_remove.append(i)
    
    # Remove paragraphs in reverse order
    for i in reversed(paragraphs_to_remove):
        try:
            doc.docx.paragraphs[i]._element.getparent().remove(doc.docx.paragraphs[i]._element)
        except Exception:
            pass
    
    # Clean empty table rows - more aggressive approach
    for table in doc.docx.tables:
        rows_to_remove = []
        for row_idx, row in enumerate(table.rows):
            # Check if all cells in the row are empty or contain only whitespace/signature lines
            all_empty_or_signature_only = True
            for cell in row.cells:
                cell_text = ' '.join([para.text for para in cell.paragraphs if para.text.strip()])
                cell_text = cell_text.strip()
                
                # Skip if cell has meaningful content (not just signature lines or whitespace)
                if cell_text and not cell_text.startswith('_________________________') and len(cell_text) > 5:
                    all_empty_or_signature_only = False
                    break
            
            if all_empty_or_signature_only:
                rows_to_remove.append(row_idx)
        
        # Remove empty rows (in reverse order to maintain indices)
        for row_idx in reversed(rows_to_remove):
            try:
                table._tbl.remove(table.rows[row_idx]._tr)
            except Exception:
                # If removal fails, just skip this row
                pass

