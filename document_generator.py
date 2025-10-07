"""
Document generation module for mail merge application.
Handles document creation, rendering, and file operations.
"""

import tempfile
import shutil
import platform
import subprocess
import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from io import BytesIO

import pandas as pd
from docx import Document
from docxtpl import DocxTemplate
from jinja2 import Template

from template_processor import extract_placeholders, clean_placeholder_elements, render_text
from data_handler import (
    create_context_from_row, calculate_totals, normalize_column_names,
    format_field_value, try_parse_date, format_date_long
)


def generate_empty_data_file(template_fields: List[str], output_path: Optional[str] = None) -> str:
    """
    Generate an empty Excel file with template fields as columns.
    
    Args:
        template_fields: List of template field names
        output_path: Optional path to save the file
        
    Returns:
        Path to the generated file
    """
    import openpyxl
    from openpyxl.styles import Alignment
    from openpyxl.worksheet.views import Selection
    
    df_empty = pd.DataFrame(columns=template_fields)
    
    if output_path is None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmpfile:
            output_path = tmpfile.name
    
    df_empty.to_excel(output_path, index=False)
    
    # Format the Excel file
    wb = openpyxl.load_workbook(output_path)
    ws = wb.active
    
    # Set column width to approximately 20 Excel characters
    for col in ws.columns:
        col_letter = col[0].column_letter
        ws.column_dimensions[col_letter].width = 20
    
    # Set alignment in each cell to left and top
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(horizontal="left", vertical="top")
    
    # Set focus/cursor on A2 when opening
    ws.sheet_view.selection = [Selection(activeCell="A2", sqref="A2")]
    wb.save(output_path)
    
    return output_path


def open_file_in_system(file_path: str) -> None:
    """
    Open a file in the system's default application.
    Only works in local environments, not in cloud deployments.
    
    Args:
        file_path: Path to the file to open
    """
    try:
        # Check if we're running in a cloud environment
        import os
        if os.getenv('STREAMLIT_CLOUD') or os.getenv('STREAMLIT_SERVER_HEADLESS'):
            # Don't try to open files in cloud environments
            return
        
        if platform.system() == "Darwin":
            subprocess.run(["open", file_path])
        elif platform.system() == "Windows":
            subprocess.run(["start", str(file_path)], shell=True)
        elif platform.system() == "Linux":
            subprocess.run(["xdg-open", str(file_path)])
    except Exception:
        # Silently fail if we can't open the file
        pass


def create_safe_filename(prefix: str, primary_value: str, secondary_value: str = "") -> str:
    """
    Create a safe filename from prefix and values.
    
    Args:
        prefix: File prefix
        primary_value: Primary value for filename
        secondary_value: Fallback value if primary is empty
        
    Returns:
        Safe filename string
    """
    value = primary_value if primary_value and not pd.isna(primary_value) else secondary_value
    safe_value = re.sub(r'[\\/*?:"<>|]', '_', str(value)).strip()
    if len(safe_value) > 80:
        safe_value = safe_value[:80]
    
    return f"{prefix} {safe_value} {pd.Timestamp.now().date()}.docx"


def render_template_preview(uploaded_template, context: Dict[str, Any]) -> str:
    """
    Render template preview using Jinja2 for display.
    
    Args:
        uploaded_template: Template file object
        context: Context dictionary with data
        
    Returns:
        Rendered text for preview
    """
    doc = Document(uploaded_template)
    preview_text = []
    
    # Paragraph preview
    for para in doc.paragraphs:
        raw_txt = para.text
        try:
            txt = Template(raw_txt).render(**context)
        except Exception:
            txt = raw_txt
        txt = render_text(txt)
        if txt.strip():
            preview_text.append(txt)
    
    # Table preview
    for tbl in doc.tables:
        raw_rows = [[cell.text for cell in row.cells] for row in tbl.rows]
        if raw_rows:
            header = raw_rows[0]
            md = "| " + " | ".join(header) + " |"
            md += "\n| " + " | ".join(["---"] * len(header)) + " |"
            
            for r in raw_rows[1:]:
                if all(not (cell or "").strip() for cell in r):
                    continue
                md += "\n| " + " | ".join([cell or "" for cell in r]) + " |"
            
            try:
                rendered_md = Template(md).render(**context)
            except Exception:
                rendered_md = md
            
            rendered_md = render_text(rendered_md)
            preview_text.append(rendered_md)
    
    return "\n\n".join(preview_text)


def generate_single_document(uploaded_template, form_data: Dict[str, str]) -> BytesIO:
    """
    Generate a single document from template and form data.
    
    Args:
        uploaded_template: Template file object
        form_data: Dictionary with form field values
        
    Returns:
        BytesIO object with the generated document
    """
    uploaded_template.seek(0)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_template:
        shutil.copyfileobj(uploaded_template, tmp_template)
        template_path = tmp_template.name
    
    doc = DocxTemplate(template_path)
    doc.render(form_data)
    
    output = BytesIO()
    doc.save(output)
    return output


def generate_documents_batch(
    uploaded_template,
    df_data: pd.DataFrame,
    field_mapping: Dict[str, str],
    square_fields: List[Dict],
    output_dir: Path,
    prefix: str,
    primary_column: str,
    secondary_column: str,
    target_lang: str = "UK"
) -> Tuple[int, List[Path]]:
    """
    Generate multiple documents from template and data.
    
    Args:
        uploaded_template: Template file object
        df_data: DataFrame with data
        field_mapping: Mapping from template fields to data columns
        square_fields: List of square bracket fields
        output_dir: Directory to save documents
        prefix: File prefix
        primary_column: Primary column for filename
        secondary_column: Secondary column for filename
        target_lang: Target language for formatting
        
    Returns:
        Tuple of (number of documents generated, list of generated file paths)
    """
    uploaded_template.seek(0)
    
    # Save uploaded template as temporary file
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_template:
        shutil.copyfileobj(uploaded_template, tmp_template)
        template_path = tmp_template.name
    
    # Calculate totals for full dataset
    total_row = calculate_totals(df_data)
    rows_all = df_data.to_dict("records")
    rows_all_with_total = rows_all + [total_row]
    
    documents_generated = 0
    generated_files = []
    
    for index, row in df_data.iterrows():
        doc = DocxTemplate(template_path)
        
        # Create context for this row
        context = create_context_from_row(row, field_mapping, square_fields, target_lang)
        
        # Add full dataset with totals for template loops
        context["rows_all"] = rows_all_with_total
        context["rows"] = context["rows_all"]  # Backward compatibility
        
        # Render the document
        doc.render(context)
        
        # Clean up placeholder elements
        clean_placeholder_elements(doc)
        
        # Create filename
        filename = create_safe_filename(
            prefix,
            row.get(primary_column, ""),
            row.get(secondary_column, "")
        )
        filepath = output_dir / filename.replace("_", " ")
        
        # Save the document
        doc.save(filepath)
        generated_files.append(filepath)
        documents_generated += 1
    
    return documents_generated, generated_files


def validate_template_before_generation(uploaded_template) -> List[str]:
    """
    Validate template before document generation.
    
    Args:
        uploaded_template: Template file object
        
    Returns:
        List of validation error messages (empty if valid)
    """
    try:
        curly_fields, _square_fields = extract_placeholders(uploaded_template)
    except Exception:
        curly_fields = []
    
    from template_processor import is_valid_jinja_var
    
    invalid_placeholders = [f for f in curly_fields if not is_valid_jinja_var(f)]
    
    if invalid_placeholders:
        return [
            "❌ Je template bevat ongeldige Jinja-plaats-houders (bijv. spaties of vreemde tekens):\n\n- " +
            "\n- ".join(invalid_placeholders) +
            "\n\n🔧 Oplossing: vervang spaties en speciale tekens door underscores. Voorbeeld: `{{ Voornaam Klant }}` → `{{ Voornaam_Klant }}`."
        ]
    
    return []
